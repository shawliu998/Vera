from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
OUTPUT = OUT_DIR / "vera-word-host-e2e-fixture.docx"


def set_run_font(run, western: str = "Calibri", east_asia: str = "Heiti SC"):
    run.font.name = western
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), western)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), western)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:cs"), east_asia)
    run._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    language = OxmlElement("w:lang")
    language.set(qn("w:val"), "en-US")
    language.set(qn("w:eastAsia"), "zh-CN")
    run._element.rPr.append(language)


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    if any(ord(character) > 127 for character in text):
        set_run_font(run, western="Heiti SC")
    else:
        set_run_font(run)
    run.font.size = Pt(11)
    return paragraph


def main():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, "2E74B5"),
        ("Heading 2", 13, 12, 6, "2E74B5"),
        ("Heading 3", 12, 8, 4, "1F4D78"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    doc.add_heading("Vera Word Add-in Host E2E Fixture", level=1)
    note = add_body(
        doc,
        "Synthetic test content only. This document contains no client, personal, or confidential information.",
    )
    note.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_heading("Tracked replacement candidate", level=2)
    add_body(
        doc,
        "The Supplier may change the Service Fees at any time by giving the Customer written notice. The revised fees take effect immediately upon notice.",
    )

    doc.add_heading("Comment candidate", level=2)
    add_body(
        doc,
        "The Customer must notify the Supplier within five business days of discovering a service issue.",
    )

    doc.add_heading("Long Chinese selection", level=2)
    add_body(
        doc,
        "供应商可在任何时间通过向客户发出书面通知调整服务费用，调整后的费用自通知发出之日起立即生效，客户不得因此解除本协议或要求退还任何已经支付的款项。为验证窄任务窗格与长中文换行，本段继续说明：客户要求至少提前三十日收到通知，并可在新费用生效前无责解除协议。",
    )

    doc.add_heading("Save and reopen marker", level=2)
    add_body(doc, "E2E-SAVE-MARKER-2026-07-21")

    doc.core_properties.title = "Vera Word Add-in Host E2E Fixture"
    doc.core_properties.subject = "Synthetic Office Add-in validation"
    doc.core_properties.author = "Vera QA"
    doc.core_properties.keywords = "synthetic, office add-in, word, e2e"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
