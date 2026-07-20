from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "fixtures" / "synthetic-software-license-review.docx"


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color="6B7280")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 12, 6),
    ("Heading 2", 13, "2E74B5", 10, 5),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    style = doc.styles[style_name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_run = header.add_run("VERA LOCAL QA  |  SYNTHETIC DOCUMENT")
set_font(header_run, size=8.5, bold=True, color="6B7280")
add_page_number(section.footer.paragraphs[0])

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(16)
title.paragraph_format.space_after = Pt(4)
title_run = title.add_run("SOFTWARE LICENCE AND DATA PROCESSING AGREEMENT")
set_font(title_run, size=23, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
subtitle_run = subtitle.add_run("Synthetic contract fixture for Vera document workflow testing")
set_font(subtitle_run, size=13, color="4B5563")

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(14)
for label, value in [
    ("Matter", "VERA-2026-0001"),
    ("Parties", "Oceanic Technology Co., Ltd. / Northstar Systems Pte. Ltd."),
    ("Status", "Unapproved test document - not legal advice"),
]:
    label_run = meta.add_run(f"{label}: ")
    set_font(label_run, bold=True)
    value_run = meta.add_run(value + "\n")
    set_font(value_run)

doc.add_heading("1. Purpose and scope", level=1)
doc.add_paragraph(
    "This synthetic agreement is created solely to test upload, extraction, citation, versioning, review, and export behaviour in the Vera local QA environment. It creates no legal rights or obligations."
)
doc.add_paragraph(
    "本合成协议仅用于测试 Vera 的文档上传、文本解析、引用定位、版本管理、人工审查及导出流程，不构成法律意见，也不产生任何真实权利义务。"
)

doc.add_heading("2. Licence grant", level=1)
doc.add_paragraph(
    "Northstar Systems grants Oceanic Technology a limited, non-exclusive, non-transferable licence to use the test software within the defined internal evaluation environment for ninety days."
)
doc.add_heading("2.1 Restrictions", level=2)
doc.add_paragraph(
    "The customer must not reverse engineer, sublicense, distribute, or use the test software to process production personal data without prior written approval."
)

doc.add_heading("3. Data processing", level=1)
doc.add_paragraph(
    "Each party must apply appropriate technical and organisational measures. Cross-border transfers require a documented transfer mechanism, a current record of processing activities, and approval by the designated privacy reviewer."
)
doc.add_paragraph(
    "任何跨境数据传输均应明确数据类别、处理目的、保存期限、接收方以及适用的传输机制；未完成审查前不得投入生产使用。"
)

doc.add_heading("4. Security incident notice", level=1)
doc.add_paragraph(
    "The processor will notify the controller without undue delay and, where practicable, within forty-eight hours after confirming a security incident affecting covered data. The notice must distinguish verified facts from preliminary conclusions."
)

doc.add_heading("5. Liability and unresolved point", level=1)
doc.add_paragraph(
    "The draft liability cap is the fees paid during the preceding twelve months. Whether confidentiality, data protection, and intellectual-property claims are excluded from the cap remains unverified and requires human review."
)

doc.add_heading("6. Approval", level=1)
doc.add_paragraph(
    "No output derived from this fixture is approved until a named reviewer records a review decision and the export status is marked approved."
)

for paragraph in doc.paragraphs:
    if any("\u4e00" <= char <= "\u9fff" for char in paragraph.text):
        for run in paragraph.runs:
            run.font.name = "Noto Sans CJK SC"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
            fonts = run._element.get_or_add_rPr().rFonts
            fonts.set(qn("w:ascii"), "Noto Sans CJK SC")
            fonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
            fonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
            fonts.set(qn("w:cs"), "Noto Sans CJK SC")
            language = OxmlElement("w:lang")
            language.set(qn("w:eastAsia"), "zh-CN")
            run._element.get_or_add_rPr().append(language)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
